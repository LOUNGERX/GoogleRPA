import os
import time
import re
from docx import Document
from playwright.sync_api import sync_playwright
import playwright_stealth

# ================= 配置区 =================
BASE_DIR = r"C:\Users\asus\Desktop\RPA_GG"
USER_DATA_DIR = os.path.join(BASE_DIR, "user_data")
SOURCE_FOLDER = r"C:\Users\asus\Desktop\test" 
OUTPUT_FOLDER = r"C:\Users\asus\Desktop\result" 
PROMPT_TEXT = "请帮我梳理成 门诊记录单 格式，不带来源编号（即去掉末尾数字）的纯净版门诊记录单,要求格式如下【基本信息】• 姓名:• 性别： • 年龄： 【主诉】【既往史】• 手术史：• 其他： 【辅助检查】【初步诊断】【处理意见】"

if not os.path.exists(OUTPUT_FOLDER): os.makedirs(OUTPUT_FOLDER)
# ==========================================

def save_to_word(content, folder_name):
    doc = Document()
    doc.add_heading(f'门诊记录单 - {folder_name}', 0)
    doc.add_paragraph(content)
    safe_name = re.sub(r'[\\/*?:"<>|]', "_", folder_name)
    save_path = os.path.join(OUTPUT_FOLDER, f"{safe_name}_汇总记录.docx")
    doc.save(save_path)
    return save_path

def run_notebooklm_final_fix():
    folder_groups = {}
    for root, dirs, files in os.walk(SOURCE_FOLDER):
        valid_files = [os.path.join(root, f) for f in files if f.lower().endswith(('.txt', '.pdf', '.docx', '.png', '.jpg', '.jpeg')) and not f.startswith('~$')]
        if valid_files: folder_groups[os.path.basename(root)] = valid_files

    with sync_playwright() as p:
        print("🚀 启动 NotebookLM 引擎...")
        context = p.chromium.launch_persistent_context(
            user_data_dir=USER_DATA_DIR, channel="msedge", headless=False,
            ignore_default_args=["--enable-automation"], args=["--disable-blink-features=AutomationControlled"]
        )
        page = context.pages[0] if context.pages else context.new_page()
        try: playwright_stealth.stealth(page)
        except: pass

        for folder_name, file_paths in folder_groups.items():
            print(f"\n📁 正在处理: 【{folder_name}】")
            try:
                page.goto("https://notebooklm.google.com/", wait_until="networkidle")
                page.get_by_text("新建笔记本").click()
                page.wait_for_url("**/notebook/*", timeout=30000)

                # --- 原始上传逻辑 ---
                print("🎯 正在上传文件...")
                with page.expect_file_chooser() as fc_info:
                    page.get_by_text("上传文件").click()
                file_chooser = fc_info.value
                file_chooser.set_files(file_paths)
                
                # 给充足的时间让“9个来源”处理完成
                print("⏳ 等待文件上传解析...")
                time.sleep(8) 

                # --- 输入与发送逻辑 (重点修复) ---
                print("⌨️ 准备输入指令...")
                # 定位输入框
                chat_box = page.locator("textarea.query-box-input, [role='textbox']").last
                chat_box.wait_for(state="visible")
                
                initial_count = page.locator("button:has-text('保存到笔记')").count()

                # 点击并模拟真人输入，触发按钮变蓝
                chat_box.click()
                chat_box.press_sequentially(PROMPT_TEXT, delay=20) 
                time.sleep(1)

                # 尝试点击蓝色发送按钮
                # 根据截图：按钮在 query-box 容器内，通常带有 mat-icon
                send_button = page.locator(".query-box button:has(mat-icon), .query-box button[aria-label*='发'], .query-box button.send-button").last
                
                if send_button.is_visible():
                    print("🚀 点击发送按钮...")
                    # 强制点击，防止被透明层遮挡
                    send_button.click(force=True, timeout=5000)
                else:
                    print("⚠️ 没找到按钮，使用 Enter 键发送...")
                    page.keyboard.press("Enter")

                # --- 原始提取逻辑 ---
                print("🤖 等待 AI 响应完成...")
                for i in range(120):
                    if page.locator("button:has-text('保存到笔记')").count() > initial_count:
                        print("✅ AI 响应已完成")
                        break
                    time.sleep(2)

                print("⏳ 预留 20 秒提取内容...")
                time.sleep(20) 

                target_locator = page.locator(".message-content")
                if target_locator.count() > 0:
                    final_text = target_locator.last.inner_text()
                    path = save_to_word(final_text, folder_name)
                    print(f"🎉 处理成功: {path}")
                else:
                    print("❌ 未能定位到回复内容")

            except Exception as e:
                print(f"❌ 错误: {e}")
                page.screenshot(path=f"fail_{folder_name}.png")

        context.close()

if __name__ == "__main__":
    run_notebooklm_final_fix()